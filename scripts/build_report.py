from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.constants import PROJECT_ROOT as ROOT
from src.io_utils import sha256_file, write_json


REPORT_PATH = ROOT / "deliverables" / "financial_performance_prediction_report.docx"
REPORT_TITLE = "财务绩效预测研究报告"
REPORT_SUBTITLE = "基于历史季度财务序列的九目标回归、泄漏安全交叉验证与 OOF 融合"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "F7F9FC"
HEADER_FILL = "D9EAF7"


TARGET_CN = {
    "Q0_TOTAL_ASSETS": "总资产",
    "Q0_TOTAL_LIABILITIES": "总负债",
    "Q0_TOTAL_STOCKHOLDERS_EQUITY": "股东权益",
    "Q0_GROSS_PROFIT": "毛利",
    "Q0_COST_OF_REVENUES": "营业成本",
    "Q0_REVENUES": "营业收入",
    "Q0_OPERATING_INCOME": "营业利润",
    "Q0_OPERATING_EXPENSES": "营业费用",
    "Q0_EBITDA": "EBITDA",
}

EXPERIMENT_CN = {
    "B0": "B0 均值预测",
    "B1": "B1 最近季度复制",
    "B2": "B2 季节性复制",
    "B3": "B3 短趋势外推",
    "B4": "B4 规则基线融合",
    "M1_ridge_history_raw": "M1 Ridge",
    "M2_hgb_history_raw": "M2 HistGBR",
    "M3a_catboost_direct_history_raw": "M3a CatBoost 直接：历史",
    "M3b_catboost_direct_history_industry": "M3b CatBoost 直接：历史+行业",
    "M3c_catboost_direct_history_metadata": "M3c CatBoost 直接：历史+元数据",
    "M3d_catboost_direct_history_metadata_engineered": "M3d CatBoost 直接：工程特征",
    "M4_catboost_residual_history_metadata_engineered": "M4 CatBoost 残差",
    "M6_oof_blend": "M6 逐目标 OOF 融合",
    "M7_accounting_postprocess": "M7 会计一致性后处理",
}


def _read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _fmt_num(value: object, digits: int = 4) -> str:
    if pd.isna(value):
        return ""
    try:
        val = float(value)
    except (TypeError, ValueError):
        return str(value)
    if abs(val) >= 1_000_000:
        return f"{val:,.0f}"
    return f"{val:.{digits}f}"


def _fmt_pct(value: object, digits: int = 2) -> str:
    if pd.isna(value):
        return ""
    return f"{float(value) * 100:.{digits}f}%"


def _short_hash(value: str, length: int = 16) -> str:
    return str(value)[:length]


def _set_run_font(
    run,
    size: float = 10.5,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = latin
    run.font.color.rgb = color or BODY
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _add_text(
    doc: Document,
    text: str,
    *,
    size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line: bool = True,
    after: float = 5,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(18)
    paragraph.paragraph_format.space_after = Pt(after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(17)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    _set_run_font(run, size=10.5)


def _set_bottom_border(paragraph, color: str = "5B9BD5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(8 if level <= 2 else 4)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _set_run_font(run, size=17, bold=True, color=ACCENT, east_asia="黑体")
    elif level == 2:
        run = paragraph.add_run(text)
        _set_run_font(run, size=13, bold=True, color=ACCENT, east_asia="黑体")
        _set_bottom_border(paragraph)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run, size=11.5, bold=True, color=BODY, east_asia="黑体")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 70, bottom: int = 70, left: int = 70, right: int = 70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _fill_cell(cell, value: object, *, bold: bool = False, size: float = 8.8, align_center: bool = True) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("" if pd.isna(value) else str(value))
    _set_run_font(run, size=size, bold=bold, east_asia="宋体")


def _add_table(
    doc: Document,
    df: pd.DataFrame,
    caption: str,
    source: str,
    *,
    font_size: float = 8.8,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(caption)
    _set_run_font(run, size=10, bold=True, color=ACCENT, east_asia="黑体")

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, col in enumerate(df.columns):
        _shade_cell(table.rows[0].cells[idx], HEADER_FILL)
        _fill_cell(table.rows[0].cells[idx], col, bold=True, size=font_size, align_center=True)
    for row_idx, (_, row) in enumerate(df.iterrows()):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row.tolist()):
            if row_idx % 2 == 1:
                _shade_cell(cells[col_idx], LIGHT_FILL)
            _fill_cell(cells[col_idx], value, size=font_size, align_center=(col_idx != 0))
    _add_source(doc, source)


def _add_source(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"资料来源：{text}")
    _set_run_font(run, size=8.5, italic=True, color=MUTED)


def _add_figure(doc: Document, path: Path, caption: str, note: str, source: str, width_cm: float = 15.2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = True
    cap_run = cap.add_run(caption)
    _set_run_font(cap_run, size=10, bold=True, color=ACCENT, east_asia="黑体")

    _add_text(doc, f"图表解读：{note}", size=9.5, color=MUTED, first_line=False, after=2)
    _add_source(doc, source)


def _add_toc(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for fld_type, instr_text in [
        ("begin", None),
        (None, ' TOC \\o "1-3" \\h \\z \\u '),
        ("separate", None),
    ]:
        run = paragraph.add_run()
        if fld_type:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), fld_type)
        else:
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = instr_text
        run._r.append(node)
    placeholder = paragraph.add_run("目录将在 Word 中打开后自动更新。")
    _set_run_font(placeholder, size=10, color=MUTED)
    end = paragraph.add_run()
    node = OxmlElement("w:fldChar")
    node.set(qn("w:fldCharType"), "end")
    end._r.append(node)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head = paragraph.add_run("第 ")
    _set_run_font(head, size=9, color=MUTED)
    run = paragraph.add_run()
    _set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    tail = paragraph.add_run(" 页")
    _set_run_font(tail, size=9, color=MUTED)


def _configure_document(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.35)
    section.bottom_margin = Cm(2.15)
    section.left_margin = Cm(2.45)
    section.right_margin = Cm(2.45)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(REPORT_TITLE)
    _set_run_font(run, size=9, color=MUTED, east_asia="宋体")
    footer = section.footer.paragraphs[0]
    _add_page_number(footer)


def _cover(doc: Document, context: dict) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_TITLE)
    _set_run_font(run, size=24, bold=True, color=ACCENT, east_asia="黑体")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(REPORT_SUBTITLE)
    _set_run_font(run, size=14, color=MUTED, east_asia="黑体")
    line = doc.add_paragraph()
    _set_bottom_border(line, color="1F4E79")

    cover_table = pd.DataFrame(
        [
            ["研究任务", "预测 Q0 最新季度 9 项财务指标"],
            ["数据规模", f"训练集 {context['train_rows']} 行 × {context['train_cols']} 列；测试集 {context['test_rows']} 行 × {context['test_cols']} 列"],
            ["验证协议", "5 折 GroupKFold；完全相同历史特征记录不跨折"],
            ["最终模型", f"{EXPERIMENT_CN.get(context['best_experiment_id'], context['best_experiment_id'])}"],
            ["OOF 平均 R²", _fmt_num(context["best_mean_oof_r2"], 4)],
            ["交付文件", "Notebook、Word/PDF 报告、submission.csv、README_delivery.md"],
            ["环境", f"Conda QuantEnv；{context['python_version']}"],
            ["报告日期", "2026 年 6 月 10 日"],
        ],
        columns=["项目", "内容"],
    )
    _add_table(doc, cover_table, "封面信息表：研究对象与交付口径", "项目自动生成结果；results/*.json。", font_size=9.5)
    doc.add_page_break()


def _toc_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("目录")
    _set_run_font(run, size=18, bold=True, color=ACCENT, east_asia="黑体")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.right_indent = Cm(1.0)
    _add_toc(paragraph)
    doc.add_page_break()


def _experiment_table(scores: pd.DataFrame) -> pd.DataFrame:
    keep = [
        "B0",
        "B1",
        "B2",
        "B3",
        "B4",
        "M1_ridge_history_raw",
        "M2_hgb_history_raw",
        "M3a_catboost_direct_history_raw",
        "M3b_catboost_direct_history_industry",
        "M3c_catboost_direct_history_metadata",
        "M3d_catboost_direct_history_metadata_engineered",
        "M4_catboost_residual_history_metadata_engineered",
        "M6_oof_blend",
    ]
    frame = scores[scores["experiment_id"].isin(keep)].copy()
    frame["实验"] = frame["experiment_id"].map(EXPERIMENT_CN).fillna(frame["experiment_id"])
    frame["OOF均值R²"] = frame["mean_r2"].map(lambda x: _fmt_num(x, 4))
    frame["分折均值标准差"] = frame["fold_mean_r2_std"].map(lambda x: _fmt_num(x, 4))
    frame["定位"] = frame["experiment_id"].map(
        {
            "B0": "指标校验",
            "B1": "最近季度规则",
            "B2": "季节性规则",
            "B3": "趋势规则",
            "B4": "规则基线",
            "M1_ridge_history_raw": "线性对照",
            "M2_hgb_history_raw": "非线性对照",
            "M3a_catboost_direct_history_raw": "直预测消融",
            "M3b_catboost_direct_history_industry": "行业消融",
            "M3c_catboost_direct_history_metadata": "元数据消融",
            "M3d_catboost_direct_history_metadata_engineered": "特征工程消融",
            "M4_catboost_residual_history_metadata_engineered": "残差主模型",
            "M6_oof_blend": "最终模型",
        }
    )
    return frame[["实验", "定位", "OOF均值R²", "分折均值标准差"]]


def _target_table(final_scores: pd.DataFrame, blend_scores: pd.DataFrame) -> pd.DataFrame:
    row = final_scores[final_scores["experiment_id"] == "M6_oof_blend"].iloc[0]
    records = []
    blend_lookup = blend_scores.set_index("target")
    for target, label in TARGET_CN.items():
        weights = blend_lookup.loc[target, "weights"]
        records.append(
            {
                "目标": label,
                "字段": target,
                "M6 OOF R²": _fmt_num(row[f"r2_{target}"], 4),
                "融合权重(B4/M3d/M4)": weights,
            }
        )
    return pd.DataFrame(records)


def _target_summary_table(target_summary: pd.DataFrame) -> pd.DataFrame:
    frame = target_summary.copy()
    frame["目标"] = frame["target"].map(TARGET_CN).fillna(frame["target"])
    frame["中位数"] = frame["median"].map(lambda x: _fmt_num(x, 0))
    frame["标准差"] = frame["std"].map(lambda x: _fmt_num(x, 0))
    frame["正值占比"] = frame["positive_rate"].map(_fmt_pct)
    frame["偏度"] = frame["skew"].map(lambda x: _fmt_num(x, 2))
    return frame[["目标", "中位数", "标准差", "正值占比", "偏度"]]


def _figure_plan() -> list[tuple[str, str, str]]:
    return [
        ("fig01_sector_distribution.png", "图 1：板块样本分布", "板块分布决定模型能否在行业维度上稳定外推。样本并非均匀覆盖，报告后续不把行业变量视为无风险增强。"),
        ("fig02_industry_top20.png", "图 2：行业 Top 20 样本数量", "行业集中度较高，行业变量能提供横截面先验，但也可能放大样本内结构偏差。"),
        ("fig03_missing_top20.png", "图 3：缺失率最高字段", "高缺失字段集中在估值、风险和远期财务项目，缺失模式本身具有信息含量。"),
        ("fig04_missing_by_quarter.png", "图 4：各历史季度平均缺失率", "Q7、Q9、Q10 的平均缺失率明显抬升，说明较早季度信息质量弱于最近季度。"),
        ("fig05_target_distributions.png", "图 5：九个目标 signed-log 分布", "目标分布厚尾且允许负值，简单 log1p 变换会破坏符号信息。"),
        ("fig06_lag_correlation_heatmap.png", "图 6：历史滞后值与 Q0 目标相关性", "最近季度和部分年度同比位置提供主要预测信息，支持 B1/B4 规则基线作为强起点。"),
        ("fig07_accounting_identity_error.png", "图 7：会计恒等式误差分布", "会计关系总体有约束力，但尾部误差较重，机械修正需要 OOF 验证。"),
        ("fig08_target_correlation_heatmap.png", "图 8：目标变量相关性", "收入、成本、利润类指标相关性较高，权益类指标更容易受重分类和异常项目影响。"),
        ("fig09_model_comparison.png", "图 9：模型 OOF 平均 R² 对比", "规则融合、残差建模和逐目标 OOF 融合依次提升，最终模型不是单一算法结果。"),
        ("fig10_target_score_heatmap.png", "图 10：各目标 OOF R² 热力图", "目标难度差异显著，股东权益是主要短板，费用、毛利和收入类目标稳定性更高。"),
        ("fig11_oof_scatter.png", "图 11：OOF 真实值与预测值散点", "signed-log 空间下大部分样本沿对角线分布，尾部公司仍是主要误差来源。"),
        ("fig12_residual_distribution.png", "图 12：OOF 残差分布", "残差分布显示多数目标误差集中，但权益和资产端残差尾部更宽。"),
        ("fig13_blend_weights.png", "图 13：逐目标 OOF 融合权重", "不同目标依赖的模型来源不同，统一权重会掩盖目标层面的结构差异。"),
    ]


def main() -> None:
    meta = _read_json(ROOT / "results" / "final_submission_manifest.json")
    env = _read_json(ROOT / "results" / "environment_audit.json")
    manifest = _read_json(ROOT / "results" / "input_manifest.json")
    best = _read_json(ROOT / "configs" / "best_config.json")
    schema = pd.read_csv(ROOT / "results" / "tables" / "schema_summary.csv")
    duplicates = pd.read_csv(ROOT / "results" / "tables" / "duplicate_summary.csv")
    target_summary = pd.read_csv(ROOT / "results" / "tables" / "target_summary.csv")
    missing_quarter = pd.read_csv(ROOT / "results" / "tables" / "missing_rate_by_quarter.csv")
    all_scores = pd.read_csv(ROOT / "results" / "tables" / "all_model_scores.csv")
    final_scores = pd.read_csv(ROOT / "results" / "tables" / "final_model_scores.csv")
    blend_scores = pd.read_csv(ROOT / "results" / "tables" / "blend_scores.csv")
    accounting_scores = pd.read_csv(ROOT / "results" / "tables" / "accounting_postprocess_scores.csv")

    train_shape = schema.set_index("dataset").loc["train"]
    test_shape = schema.set_index("dataset").loc["test"]
    context = {
        "train_rows": int(train_shape["rows"]),
        "train_cols": int(train_shape["cols"]),
        "test_rows": int(test_shape["rows"]),
        "test_cols": int(test_shape["cols"]),
        "best_experiment_id": best["best_experiment_id"],
        "best_mean_oof_r2": best["best_mean_oof_r2"],
        "python_version": env["python_version"],
    }

    doc = Document()
    _configure_document(doc)
    _cover(doc, context)
    _toc_page(doc)

    _add_heading(doc, "摘要与核心结论", 1)
    _add_text(
        doc,
        "本项目是九目标监督回归任务：以公司过去 10 个季度的财务序列和元数据预测 Q0 最新季度的 9 项财务指标。"
        "评估口径为 9 个目标 R² 的算术平均值。报告不使用排行榜或测试集反馈选模，所有模型、融合权重和后处理选择均来自 GroupKFold OOF 结果。",
    )
    for bullet in [
        f"最终模型为 {EXPERIMENT_CN[best['best_experiment_id']]}，OOF 平均 R² 为 {_fmt_num(best['best_mean_oof_r2'], 4)}，分折均值标准差为 {_fmt_num(final_scores.loc[final_scores['experiment_id'] == 'M6_oof_blend', 'fold_mean_r2_std'].iloc[0], 4)}。",
        "规则基线并不弱。B4 规则融合 OOF 平均 R² 为 0.7831，说明最近季度复制、季节性复制和短趋势外推已经捕捉了大部分财务序列惯性。",
        "CatBoost 直接预测的边际收益主要来自元数据，M3a/M3b/M3c/M3d 的 OOF 平均 R² 分别为 0.5543、0.5695、0.6084 和 0.6090；人工序列特征对直接模型的增益有限。",
        "残差模型是核心改进。M4 在 B4 基线残差上训练，OOF 平均 R² 达到 0.8253，说明模型更擅长学习偏离历史惯性的部分，而不是直接重建公司规模。",
        "会计一致性后处理没有被强行采用。OOF 检验显示 none 的平均 R² 最高，强制资产负债表或利润表恒等式会降低预测精度。",
        "最困难目标是股东权益，M6 OOF R² 为 0.4009；毛利、收入、营业费用和 EBITDA 等经营类目标均超过 0.96。",
    ]:
        _add_bullet(doc, bullet)

    summary = pd.DataFrame(
        [
            ["最终模型", EXPERIMENT_CN[best["best_experiment_id"]]],
            ["OOF 平均 R²", _fmt_num(best["best_mean_oof_r2"], 4)],
            ["最强规则基线", "B4，OOF 平均 R² 0.7831"],
            ["最强单模型", "M4 残差模型，OOF 平均 R² 0.8253"],
            ["会计后处理", "OOF 选择 none，不强制修正"],
            ["submission", f"406 行 × 10 列；SHA256 {_short_hash(meta['submission_sha256'])}..."],
        ],
        columns=["项目", "结论"],
    )
    _add_table(doc, summary, "表 1：核心结论摘要", "results/tables/final_model_scores.csv、results/final_submission_manifest.json。", font_size=9.5)

    _add_heading(doc, "1. 任务理解、数据结构与冻结", 1)
    _add_heading(doc, "1.1 任务口径", 2)
    _add_text(
        doc,
        "目标列为总资产、总负债、股东权益、毛利、营业成本、营业收入、营业利润、营业费用和 EBITDA。"
        "时间顺序为 Q10 到 Q1，再预测 Q0。Id 只作为记录标识，不进入模型特征。"
        "最终 submission 的列顺序完全沿用 sample_submission.csv。"
    )
    schema_display = schema.copy()
    schema_display["缺失率"] = schema_display["missing_rate"].map(_fmt_pct)
    schema_display = schema_display.rename(columns={"dataset": "数据集", "rows": "行数", "cols": "列数", "id_unique": "Id唯一", "feature_cols": "特征列", "target_cols": "目标列"})
    _add_table(doc, schema_display[["数据集", "行数", "列数", "Id唯一", "缺失率", "特征列", "目标列"]], "表 2：数据集 schema 审计", "results/tables/schema_summary.csv。")

    raw = pd.DataFrame(list(manifest["raw_files"].values()))[["name", "size_bytes", "sha256"]]
    raw["大小(MB)"] = raw["size_bytes"].map(lambda x: f"{x / 1024 / 1024:.2f}")
    raw["SHA256前16位"] = raw["sha256"].map(_short_hash)
    raw = raw.rename(columns={"name": "原始文件"})
    _add_table(doc, raw[["原始文件", "大小(MB)", "SHA256前16位"]], "表 3：原始文件冻结清单", "results/input_manifest.json；完整哈希见项目 manifest。")

    _add_heading(doc, "1.2 数据质量和泄漏风险", 2)
    _add_text(
        doc,
        "训练集和测试集均存在完全相同的历史特征记录。若使用普通 KFold，相同财务历史可能同时进入训练折和验证折，验证分数会被高估。"
        "因此全部 OOF 实验使用基于稳定特征哈希的 GroupKFold，并把重复特征记录固定在同一折。"
    )
    dup = duplicates.rename(
        columns={
            "scope": "范围",
            "duplicate_rows": "重复行",
            "duplicate_groups": "重复组",
            "max_group_size": "最大组大小",
            "notes": "说明",
        }
    )
    _add_table(doc, dup, "表 4：重复样本与跨集特征匹配审计", "results/tables/duplicate_summary.csv。", font_size=8.2)

    mq = missing_quarter.copy()
    mq["平均缺失率"] = mq["avg_missing_rate"].map(_fmt_pct)
    mq["中位缺失率"] = mq["median_missing_rate"].map(_fmt_pct)
    mq = mq.rename(columns={"quarter": "季度", "observed_columns": "字段数"})
    _add_table(doc, mq[["季度", "平均缺失率", "中位缺失率", "字段数"]], "表 5：季度维度缺失率", "results/tables/missing_rate_by_quarter.csv。")

    _add_heading(doc, "2. EDA：样本分布、缺失模式与目标特征", 1)
    _add_text(
        doc,
        "EDA 的目的不是展示图表数量，而是回答建模前的三个问题：样本结构是否稳定、缺失是否有时间规律、目标是否适合直接线性外推。"
        "结论是：行业和板块分布不均、较早季度缺失率更高，且目标分布厚尾明显。"
    )
    for idx, (figure, caption, note) in enumerate(_figure_plan()[:8], start=1):
        _add_figure(doc, ROOT / "figures" / figure, caption, note, "results/tables/*.csv 与 figures/*.png 自动生成。")

    target_display = _target_summary_table(target_summary)
    _add_table(doc, target_display, "表 6：目标变量分布特征", "results/tables/target_summary.csv。", font_size=8.6)
    _add_text(
        doc,
        "目标分布的一个关键特征是厚尾和符号混合。营业利润正值占比约 48.40%，EBITDA 正值占比约 45.38%，说明部分目标天然跨零。"
        "这也是报告采用 signed-log 可视化、但最终模型保留原值预测的重要原因。"
    )

    _add_heading(doc, "3. 建模路线：先规则基线，再机器学习", 1)
    _add_heading(doc, "3.1 规则基线的解释价值", 2)
    _add_text(
        doc,
        "财务报表具有明显的序列惯性。若最近季度复制已经很强，机器学习模型必须证明自己能够学习额外信息；否则复杂模型只是重述最近一期。"
        "B4 逐目标融合 B1、B2 和 B3，OOF 平均 R² 达到 0.7831，是后续模型和报告叙事的基准线。"
    )
    exp_table = _experiment_table(all_scores)
    _add_table(doc, exp_table, "表 7：主要实验 OOF 结果汇总", "results/tables/all_model_scores.csv、results/experiment_log.csv。", font_size=8.2)

    _add_heading(doc, "3.2 元数据消融和残差建模", 2)
    _add_text(
        doc,
        "CatBoost 直接模型的消融结果显示，行业分类只带来小幅改善，完整元数据的贡献更明显；但工程化序列特征对直接预测的均值提升很小。"
        "原因在于九个目标的量级差异较大，直接模型需要同时处理公司规模、行业结构和时间趋势，任务负担较重。"
    )
    _add_text(
        doc,
        "残差建模把任务拆为两层：先由 B4 解释稳定的历史惯性，再由 CatBoost 学习偏离惯性的部分。"
        "这种拆法降低了模型直接学习公司绝对规模的压力，M4 平均 R² 从 M3d 的 0.6090 提升到 0.8253。"
    )
    for idx, (figure, caption, note) in enumerate(_figure_plan()[8:10], start=9):
        _add_figure(doc, ROOT / "figures" / figure, caption, note, "results/tables/all_model_scores.csv。")

    _add_heading(doc, "4. 最终融合：逐目标选择模型来源", 1)
    _add_text(
        doc,
        "最终模型不使用统一权重。每个目标在 OOF 上独立搜索 B4、M3d 和 M4 的融合权重，目标函数为该目标的 R²。"
        "这一设计符合任务评价方式：先逐目标计算 R²，再取算术平均值。"
    )
    _add_table(doc, _target_table(final_scores, blend_scores), "表 8：最终模型各目标 R² 与融合权重", "results/tables/blend_scores.csv、results/tables/final_model_scores.csv。", font_size=8.1)
    _add_text(
        doc,
        "融合权重揭示了目标结构差异。总负债、毛利、营业成本、营业利润和 EBITDA 完全依赖残差模型；股东权益完全依赖 M3d 直接模型；"
        "营业费用主要沿用 B4 规则基线。这说明不同财务项目的信息来源不同，统一建模路线会损失目标层面的稳定性。"
    )
    for idx, (figure, caption, note) in enumerate(_figure_plan()[10:13], start=11):
        _add_figure(doc, ROOT / "figures" / figure, caption, note, "results/oof/m6_oof_blend.csv、results/tables/blend_scores.csv。")

    _add_heading(doc, "5. 会计一致性后处理：为什么选择不修正", 1)
    _add_text(
        doc,
        "资产负债表和利润表恒等式是重要的财务约束，但预测任务的评价指标是逐目标 R²。"
        "若强制修正能提升报表逻辑却显著降低 OOF R²，则不应进入最终提交。"
    )
    acct = accounting_scores.copy()
    acct["平均R²"] = acct["mean_r2"].map(lambda x: _fmt_num(x, 4))
    acct = acct.rename(columns={"adjustment": "后处理方案"})
    _add_table(doc, acct[["后处理方案", "平均R²"]], "表 9：会计一致性后处理 OOF 检验", "results/tables/accounting_postprocess_scores.csv。")
    _add_text(
        doc,
        "结果显示，none 的平均 R² 为 0.8572，是所有方案中最高。强制 balance_sheet_equity 会把股东权益 R² 从 0.4009 降至 0.2165；"
        "强制 operating_income_identity 则会使营业利润 R² 大幅转负。最终提交不做机械会计修正。"
    )

    _add_heading(doc, "6. 可复现性、交付验证与风险提示", 1)
    _add_heading(doc, "6.1 交付验证", 2)
    validation = _read_json(ROOT / "results" / "delivery_validation.json")
    delivery = pd.DataFrame(
        [
            ["Notebook", "deliverables/financial_performance_prediction_final.ipynb", _short_hash(validation["notebook_sha256"])],
            ["Word 报告", "deliverables/financial_performance_prediction_report.docx", "本次构建后写入 report_manifest"],
            ["PDF 预览", "deliverables/financial_performance_prediction_report.pdf", "本次导出后写入 report_manifest"],
            ["submission", "deliverables/submission.csv", _short_hash(validation["final_submission_sha256"])],
            ["验收状态", "scripts/validate_delivery.py", "PASS" if validation["passed"] else "FAIL"],
        ],
        columns=["交付项", "路径", "校验"],
    )
    _add_table(doc, delivery, "表 10：交付物和验证状态", "results/delivery_validation.json、results/package_manifest.json。", font_size=8.5)
    _add_text(
        doc,
        "全部自动化命令使用 QuantEnv。最终验收检查覆盖原始文件哈希、submission 行列和列顺序、Notebook 执行、DOCX/PDF 存在性和实验日志完整性。"
    )

    _add_heading(doc, "6.2 风险提示", 2)
    for bullet in [
        "样本量只有 1,624 条训练记录，九个目标均存在厚尾样本；高 R² 不能等同于对所有尾部公司的精确预测。",
        "股东权益目标仍是短板，M6 OOF R² 为 0.4009。该目标可能受重分类、非经常性项目和会计调整影响，历史序列解释力弱于经营类指标。",
        "元数据来自同一数据文件且 train/test 均提供，允许用于主模型；但若未来应用场景无法在预测时获得同等元数据，模型表现需要重新评估。",
        "LightGBM、XGBoost 和 Optuna 未用于最终路线。该选择是为了避免无必要地改动长期使用的 QuantEnv，不代表这些方法无价值。",
        "会计恒等式没有强制写入 submission。报告已在 OOF 上验证强制修正会降低目标 R²，最终以比赛评价指标和可复现验证为准。",
    ]:
        _add_bullet(doc, bullet)

    _add_heading(doc, "7. 结论", 1)
    _add_text(
        doc,
        "本项目的有效路线可以概括为：用财务序列规则基线解释稳定惯性，用残差模型学习偏离惯性的部分，再按目标在 OOF 上融合。"
        "该路线同时满足课程要求中的 EDA、分析过程、机器学习训练、交叉验证和 submission 生成。"
    )
    _add_text(
        doc,
        "从结果看，最终模型的优势并非来自单一算法，而来自验证协议和任务拆解：GroupKFold 控制重复样本泄漏，B4 提供强规则基线，M4 提升残差解释，M6 再按目标选择最合适的信息来源。"
        "在不使用测试集反馈和不手工调整报告数字的前提下，M6 是当前最稳健的交付版本。"
    )

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))

    report_manifest = {
        "report_title": REPORT_TITLE,
        "language": "zh-CN",
        "style_reference": "D:/PyCharm/Quant/consume_bonus",
        "best_experiment_id": meta["final_experiment_id"],
        "figure_files": [name for name, _, _ in _figure_plan()],
        "table_files": [
            "schema_summary.csv",
            "duplicate_summary.csv",
            "missing_rate_by_quarter.csv",
            "target_summary.csv",
            "all_model_scores.csv",
            "final_model_scores.csv",
            "blend_scores.csv",
            "accounting_postprocess_scores.csv",
        ],
        "report_sha256": sha256_file(REPORT_PATH),
        "submission_sha256": meta["submission_sha256"],
        "generated_utc": datetime.now(timezone.utc).isoformat(),
    }
    write_json(ROOT / "results" / "report_manifest.json", report_manifest)
    print(f"Saved {REPORT_PATH}")


if __name__ == "__main__":
    main()
