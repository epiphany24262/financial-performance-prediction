from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path

import nbformat as nbf
import pandas as pd
from docx import Document
from nbconvert.preprocessors import ExecutePreprocessor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.constants import ID_COLUMN, PROJECT_ROOT as ROOT, RAW_FILES, SUBMISSION_COLUMNS
from src.io_utils import sha256_file, write_json


REQUIRED_DELIVERABLES = {
    "notebook": ROOT / "deliverables" / "financial_performance_prediction_final.ipynb",
    "report": ROOT / "deliverables" / "financial_performance_prediction_report.docx",
    "report_pdf": ROOT / "deliverables" / "financial_performance_prediction_report.pdf",
    "submission": ROOT / "deliverables" / "submission.csv",
    "readme": ROOT / "deliverables" / "README_delivery.md",
}


def _read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _check(condition: bool, name: str, details: str, checks: list[dict]) -> None:
    checks.append({"name": name, "passed": bool(condition), "details": details})
    if not condition:
        raise AssertionError(f"{name}: {details}")


def _execute_notebook(source: Path, kernel_name: str = "quantenv") -> None:
    nb = nbf.read(str(source), as_version=4)
    ep = ExecutePreprocessor(timeout=1800, kernel_name=kernel_name)
    ep.preprocess(nb, {"metadata": {"path": str(ROOT)}})


def main() -> None:
    checks: list[dict] = []

    input_manifest = _read_json(ROOT / "results" / "input_manifest.json")
    env = _read_json(ROOT / "results" / "environment_audit.json")
    final_manifest = _read_json(ROOT / "results" / "final_submission_manifest.json")
    report_manifest = _read_json(ROOT / "results" / "report_manifest.json")
    report_pdf_validation = _read_json(ROOT / "results" / "report_pdf_validation.json")
    notebook_manifest = _read_json(ROOT / "results" / "notebook_manifest.json")
    package_manifest = _read_json(ROOT / "results" / "package_manifest.json")
    best_config = _read_json(ROOT / "configs" / "best_config.json")

    _check(env["conda_environment_name"] == "QuantEnv", "environment_name", "Environment must be QuantEnv", checks)
    _check(bool(env["active_path_hint_contains_quantenv"]), "environment_path", "Python executable should point to QuantEnv", checks)

    for item in input_manifest["raw_files"].values():
        raw_path = ROOT / item["name"]
        _check(raw_path.exists(), f"raw_exists:{item['name']}", "Raw file must exist", checks)
        _check(sha256_file(raw_path) == item["sha256"], f"raw_hash:{item['name']}", "Raw file hash must match frozen manifest", checks)

    submission = pd.read_csv(REQUIRED_DELIVERABLES["submission"])
    sample = pd.read_csv(ROOT / "sample_submission.csv")
    _check(tuple(submission.shape) == tuple(sample.shape), "submission_shape", "Submission shape must match sample template", checks)
    _check(list(submission.columns) == SUBMISSION_COLUMNS, "submission_columns", "Submission columns must preserve template order", checks)
    _check(submission[ID_COLUMN].equals(sample[ID_COLUMN]), "submission_ids", "Submission Id column must match sample submission exactly", checks)
    values = submission.drop(columns=[ID_COLUMN])
    _check(not values.isna().any().any(), "submission_nan", "Submission must not contain NaN", checks)
    _check(not pd.DataFrame(values).isin([float("inf"), float("-inf")]).any().any(), "submission_inf", "Submission must not contain inf", checks)
    _check(sha256_file(REQUIRED_DELIVERABLES["submission"]) == final_manifest["submission_sha256"], "submission_sha", "Submission sha must match manifest", checks)

    notebook_path = REQUIRED_DELIVERABLES["notebook"]
    report_path = REQUIRED_DELIVERABLES["report"]
    report_pdf_path = REQUIRED_DELIVERABLES["report_pdf"]
    readme_path = REQUIRED_DELIVERABLES["readme"]
    _check(notebook_path.exists(), "notebook_exists", "Executed notebook must exist", checks)
    _check(report_path.exists(), "report_exists", "Report DOCX must exist", checks)
    _check(report_pdf_path.exists(), "report_pdf_exists", "Report PDF preview must exist", checks)
    _check(readme_path.exists(), "readme_exists", "README must exist", checks)

    _check(sha256_file(notebook_path) == notebook_manifest["output_sha256"], "notebook_sha", "Notebook sha must match manifest", checks)
    _check(sha256_file(report_path) == report_manifest["report_sha256"], "report_sha", "Report sha must match manifest", checks)
    _check(sha256_file(report_pdf_path) == report_manifest["report_pdf_sha256"], "report_pdf_sha", "Report PDF sha must match report manifest", checks)
    _check(sha256_file(report_pdf_path) == report_pdf_validation["pdf_sha256"], "report_pdf_validation_sha", "Report PDF sha must match PDF validation", checks)
    _check(report_pdf_validation["page_count"] > 0, "report_pdf_pages", "Report PDF must have pages", checks)
    _check(bool(report_pdf_validation["first_page_has_title"]), "report_pdf_title", "Report PDF first page must contain title text", checks)
    for preview in report_pdf_validation["preview_images"]:
        _check((ROOT / preview).exists(), f"report_pdf_preview:{preview}", "PDF preview image must exist", checks)
    _check(sha256_file(readme_path) == package_manifest["readme_sha256"], "readme_sha", "README sha must match package manifest", checks)
    _check(best_config["best_experiment_id"] in {"M6_oof_blend", "M7_accounting_postprocess"}, "best_config", "Best config should reflect final blend stage", checks)

    report_doc = Document(str(report_path))
    _check(len(report_doc.paragraphs) > 10, "report_paragraphs", "Report should contain multiple paragraphs", checks)
    _check(len(report_doc.tables) >= 5, "report_tables", "Report should contain multiple tables", checks)
    _check(len(report_doc.inline_shapes) >= 8, "report_figures", "Report should embed the key figures", checks)
    report_text_head = "\n".join(paragraph.text for paragraph in report_doc.paragraphs[:20])
    _check(
        "财务绩效预测研究报告" in report_text_head
        or "Financial Performance Prediction Report" in report_text_head,
        "report_title",
        "Report title should be present",
        checks,
    )

    _execute_notebook(ROOT / "notebooks" / "financial_performance_prediction_final.ipynb", kernel_name=notebook_manifest["kernel_name"])

    experiment_log = pd.read_csv(ROOT / "results" / "experiment_log.csv")
    expected_experiments = {
        "B0",
        "B1",
        "B2",
        "B3",
        "B4",
        "M1_ridge_history_raw",
        "M2_hgb_history_raw",
        "M3a_catboost_direct_history_raw",
        "M3b_catboost_direct_history_industry",
        "M3c_catboost_direct_history_metadata",
        "M3d_catboost_direct_history_metadata_engineered",
        "M4_catboost_residual_history_metadata_engineered",
        "M6_oof_blend",
        "M7_accounting_postprocess",
    }
    _check(
        len(experiment_log) >= len(expected_experiments) * 45,
        "experiment_log_rows",
        "Experiment log should include 5 folds x 9 targets for each required experiment",
        checks,
    )
    _check(expected_experiments.issubset(set(experiment_log["experiment_id"])), "experiment_log_ids", "Required experiments must be logged", checks)

    validation = {
        "passed": True,
        "timestamp_utc": datetime.now(timezone.utc).isoformat(),
        "checks": checks,
        "best_experiment_id": best_config["best_experiment_id"],
        "best_mean_oof_r2": best_config["best_mean_oof_r2"],
        "final_submission_sha256": final_manifest["submission_sha256"],
        "report_sha256": report_manifest["report_sha256"],
        "report_pdf_sha256": report_manifest["report_pdf_sha256"],
        "notebook_sha256": notebook_manifest["output_sha256"],
    }
    write_json(ROOT / "results" / "delivery_validation.json", validation)
    print(json.dumps(validation, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
