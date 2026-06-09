from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def file_metadata(path: Path) -> dict[str, Any]:
    stat = path.stat()
    return {
        "path": str(path),
        "name": path.name,
        "size_bytes": stat.st_size,
        "mtime_utc": pd.to_datetime(stat.st_mtime, unit="s", utc=True).isoformat(),
        "sha256": sha256_file(path),
    }


def read_docx_text(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return {
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraph_preview": paragraphs[:12],
        "text_preview": "\n".join(paragraphs[:20]),
    }


def read_text_file(path: Path, max_chars: int = 5000) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return {
        "char_count": len(text),
        "line_count": text.count("\n") + (1 if text else 0),
        "preview": text[:max_chars],
    }


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def write_json(path: Path, payload: Any) -> None:
    ensure_parent(path)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_csv(path: Path, df: pd.DataFrame) -> None:
    ensure_parent(path)
    df.to_csv(path, index=False)


def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path)

